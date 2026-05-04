"""Multi-format document extraction (PDF, Word, spreadsheets, text, markdown)."""

from __future__ import annotations

import csv
import io
import re
import tempfile
from pathlib import Path

import fitz  # pymupdf
import pandas as pd
import pdfplumber
from docx import Document as DocxDocument
from pypdf import PdfReader


def _sanitize(s: str) -> str:
    return re.sub(r"\s+", " ", s or "").strip()


def load_pdf(path: Path) -> list[dict]:
    sections: list[dict] = []
    text_pages: list[str] = []
    with fitz.open(path) as doc:
        for i in range(len(doc)):
            page = doc.load_page(i)
            text_pages.append(_sanitize(page.get_text("text")))

    full = "\n\n".join(t for t in text_pages if t)
    if full:
        sections.append(
            {
                "text": full,
                "section_path": "",
                "section": "",
                "start_page": 1,
                "end_page": len(text_pages),
                "metadata": {"extraction": "pymupdf_text"},
            }
        )

    table_fragments: list[str] = []
    with pdfplumber.open(path) as pdf:
        for i, page in enumerate(pdf.pages):
            try:
                tables = page.extract_tables() or []
            except Exception:
                tables = []
            for ti, table in enumerate(tables):
                if not table:
                    continue
                lines = []
                for row in table:
                    cells = [str(c or "").strip() for c in row]
                    if any(cells):
                        lines.append(" | ".join(cells))
                if lines:
                    table_fragments.append(f"--- Page {i + 1} Table {ti + 1} ---\n" + "\n".join(lines))
    if table_fragments:
        sections.append(
            {
                "text": "\n\n".join(table_fragments),
                "section_path": "tables",
                "section": "Extracted tables",
                "start_page": 1,
                "end_page": len(text_pages) or 1,
                "metadata": {"extraction": "pdfplumber_tables"},
            }
        )

    if not sections:
        reader = PdfReader(str(path))
        buf = []
        for page in reader.pages:
            buf.append(_sanitize(page.extract_text() or ""))
        joined = "\n\n".join(buf)
        if joined:
            sections.append(
                {
                    "text": joined,
                    "section_path": "",
                    "section": "",
                    "start_page": 1,
                    "end_page": len(reader.pages),
                    "metadata": {"extraction": "pypdf_fallback"},
                }
            )
    return sections


def load_docx(path: Path) -> list[dict]:
    doc = DocxDocument(str(path))
    parts: list[str] = []
    for p in doc.paragraphs:
        t = _sanitize(p.text)
        if t:
            parts.append(t)
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [_sanitize(c.text) for c in row.cells]
            if any(cells):
                rows.append(" | ".join(cells))
        if rows:
            parts.append("--- Table ---\n" + "\n".join(rows))
    text = "\n\n".join(parts)
    if not text:
        return []
    return [
        {
            "text": text,
            "section_path": "body",
            "section": path.stem,
            "start_page": 0,
            "end_page": 0,
            "metadata": {"extraction": "python-docx"},
        }
    ]


def load_tabular(path: Path, raw: bytes) -> list[dict]:
    ext = path.suffix.lower()
    sections: list[dict] = []
    if ext == ".csv":
        decoded = raw.decode("utf-8", errors="replace")
        reader = csv.reader(io.StringIO(decoded))
        rows = list(reader)
        if not rows:
            return []
        header = rows[0]
        lines = [" | ".join(header)]
        for row in rows[1:10100]:
            lines.append(" | ".join(row))
        sections.append(
            {
                "text": "\n".join(lines),
                "section_path": "csv",
                "section": path.name,
                "start_page": 0,
                "end_page": 0,
                "metadata": {"rows": str(len(rows)), "extraction": "csv"},
            }
        )
        return sections

    if ext in {".xlsx", ".xlsm"}:
        xls = pd.ExcelFile(io.BytesIO(raw))
        for sheet in xls.sheet_names:
            df = pd.read_excel(xls, sheet_name=sheet, header=None)
            df = df.fillna("")
            sample = df.head(5000)
            text = sample.to_csv(index=False, sep="|")
            sections.append(
                {
                    "text": _sanitize(text),
                    "section_path": f"sheet:{sheet}",
                    "section": sheet,
                    "start_page": 0,
                    "end_page": 0,
                    "metadata": {"extraction": "pandas_xlsx", "sheet": sheet},
                }
            )
        return sections

    if ext == ".xls":
        df = pd.read_excel(io.BytesIO(raw), header=None)
        df = df.fillna("")
        text = df.head(5000).to_csv(index=False, sep="|")
        sections.append(
            {
                "text": _sanitize(text),
                "section_path": "sheet",
                "section": path.stem,
                "start_page": 0,
                "end_page": 0,
                "metadata": {"extraction": "pandas_xls"},
            }
        )
        return sections

    return []


def load_plaintext(path: Path, raw: bytes) -> list[dict]:
    text = raw.decode("utf-8", errors="replace")
    return [
        {
            "text": text,
            "section_path": "",
            "section": path.stem,
            "start_page": 0,
            "end_page": 0,
            "metadata": {"extraction": "utf-8"},
        }
    ]


def load_document(filename: str, raw: bytes) -> list[dict]:
    path = Path(filename)
    ext = path.suffix.lower()
    if ext == ".pdf":
        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as f:
            f.write(raw)
            tmp = Path(f.name)
        try:
            return load_pdf(tmp)
        finally:
            tmp.unlink(missing_ok=True)
    if ext == ".docx":
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            f.write(raw)
            tmp = Path(f.name)
        try:
            return load_docx(tmp)
        finally:
            tmp.unlink(missing_ok=True)
    if ext in {".csv", ".xlsx", ".xls", ".xlsm"}:
        return load_tabular(path, raw)
    if ext in {".txt", ".md", ".json", ".yaml", ".yml", ".html", ".htm"}:
        return load_plaintext(path, raw)
    raise ValueError(f"Unsupported file type: {ext or 'unknown'}")
